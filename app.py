import streamlit as st
from docx import Document
from io import BytesIO

def extract_paragraph_style(paragraph):
    style = {
        "font_name": paragraph.style.font.name,
        "font_size": paragraph.style.font.size.pt if paragraph.style.font.size else 11,
        "bold": paragraph.style.font.bold,
        "italic": paragraph.style.font.italic,
        "alignment": paragraph.alignment
    }
    return style

def apply_style_to_paragraph(paragraph, style):
    run = paragraph.add_run()
    font = run.font
    font.name = style["font_name"]
    font.size = style["font_size"]
    font.bold = style["bold"]
    font.italic = style["italic"]
    paragraph.alignment = style["alignment"]
    return run

def main():
    st.title("📝 AI Word Template Formatter")

    template_file = st.file_uploader("Upload Template Word Document (.docx)", type=["docx"])
    content_file = st.file_uploader("Upload New Content (Text or Word)", type=["docx", "txt"])

    if template_file and content_file:
        # Extract style from first paragraph of template
        template_doc = Document(template_file)
        if not template_doc.paragraphs:
            st.error("Template document has no paragraphs.")
            return
        style = extract_paragraph_style(template_doc.paragraphs[0])

        # Load content
        if content_file.name.endswith(".txt"):
            raw_text = content_file.read().decode("utf-8")
            content_lines = raw_text.strip().split("\n")
        else:
            new_doc = Document(content_file)
            content_lines = [p.text for p in new_doc.paragraphs if p.text.strip()]

        # Create new formatted document
        final_doc = Document()
        for line in content_lines:
            p = final_doc.add_paragraph()
            run = apply_style_to_paragraph(p, style)
            run.text = line

        # Save to buffer
        buffer = BytesIO()
        final_doc.save(buffer)
        buffer.seek(0)

        st.success("✅ Document formatted successfully!")
        st.download_button("📥 Download Formatted Word Document", buffer, file_name="formatted.docx")

if __name__ == "__main__":
    main()
